"""
Tạo file Word (.docx) cho bản tin tài chính cuối ngày.
"""
import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CATEGORY_CONFIG = {
    "vi_mo_viet_nam": {
        "label": "VĨ MÔ VIỆT NAM",
        "icon": "🏛",
        "color": RGBColor(0x1A, 0x52, 0x76),
    },
    "thi_truong": {
        "label": "THỊ TRƯỜNG",
        "icon": "📈",
        "color": RGBColor(0x1E, 0x84, 0x49),
    },
    "the_gioi": {
        "label": "THẾ GIỚI",
        "icon": "🌏",
        "color": RGBColor(0x78, 0x42, 0x12),
    },
    "doanh_nghiep": {
        "label": "TIN NỔI BẬT TỪ DOANH NGHIỆP",
        "icon": "🏢",
        "color": RGBColor(0x51, 0x2E, 0x5F),
    },
}


def _add_horizontal_rule(doc):
    """Vẽ đường kẻ ngang ngăn cách section."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)


def _set_font(run, bold=False, size=11, color=None, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def build_docx(summary: dict, session: str) -> bytes:
    """
    Tạo file Word và trả về bytes.
    session: 'Sáng' hoặc 'Chiều'
    """
    doc = Document()

    # ── Thiết lập lề trang ────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Header: Tiêu đề chính ──────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run = title_p.add_run("BẢN TIN TÀI CHÍNH")
    _set_font(run, bold=True, size=22, color=RGBColor(0x1A, 0x3A, 0x6B))

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_before = Pt(0)
    subtitle_p.paragraph_format.space_after = Pt(2)
    now = datetime.now()
    date_str = now.strftime("%d/%m/%Y")
    time_str = now.strftime("%H:%M")
    run2 = subtitle_p.add_run(f"Bản tin {session}  •  {date_str}  •  {time_str}")
    _set_font(run2, size=11, color=RGBColor(0x77, 0x77, 0x77))

    _add_horizontal_rule(doc)

    total = sum(len(summary.get(k, [])) for k in CATEGORY_CONFIG)
    intro_p = doc.add_paragraph()
    intro_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro_p.paragraph_format.space_after = Pt(12)
    run3 = intro_p.add_run(f"Tổng hợp {total} tin nổi bật từ 24hmoney · cafef · vietstock · baochinhphu")
    _set_font(run3, size=10, color=RGBColor(0x99, 0x99, 0x99), italic=True)

    # ── Từng chuyên mục ────────────────────────────────────────────────────────
    for key, cfg in CATEGORY_CONFIG.items():
        items = summary.get(key, [])
        if not items:
            continue

        # Tiêu đề chuyên mục
        cat_p = doc.add_paragraph()
        cat_p.paragraph_format.space_before = Pt(14)
        cat_p.paragraph_format.space_after = Pt(6)
        run_icon = cat_p.add_run(f"{cfg['icon']}  {cfg['label']}")
        _set_font(run_icon, bold=True, size=14, color=cfg["color"])

        _add_horizontal_rule(doc)

        for idx, item in enumerate(items, 1):
            title = item.get("title", "")
            summary_text = item.get("summary", "")
            key_points = item.get("key_points", [])
            impact = item.get("impact", "")
            source = item.get("source", "")
            url = item.get("url", "")

            # Tiêu đề bài báo
            art_title_p = doc.add_paragraph()
            art_title_p.paragraph_format.space_before = Pt(10)
            art_title_p.paragraph_format.space_after = Pt(3)
            run_num = art_title_p.add_run(f"{idx}.  ")
            _set_font(run_num, bold=True, size=12, color=cfg["color"])
            run_title = art_title_p.add_run(title)
            _set_font(run_title, bold=True, size=12)

            # Đoạn tóm tắt
            if summary_text:
                sum_p = doc.add_paragraph()
                sum_p.paragraph_format.space_before = Pt(2)
                sum_p.paragraph_format.space_after = Pt(4)
                sum_p.paragraph_format.left_indent = Cm(0.7)
                run_sum = sum_p.add_run(summary_text)
                _set_font(run_sum, size=11)

            # Key points (bullet points)
            for point in key_points:
                kp_p = doc.add_paragraph(style="List Bullet")
                kp_p.paragraph_format.left_indent = Cm(1.2)
                kp_p.paragraph_format.space_before = Pt(1)
                kp_p.paragraph_format.space_after = Pt(1)
                run_kp = kp_p.add_run(point)
                _set_font(run_kp, size=10.5)

            # Tác động
            if impact:
                imp_p = doc.add_paragraph()
                imp_p.paragraph_format.space_before = Pt(4)
                imp_p.paragraph_format.space_after = Pt(2)
                imp_p.paragraph_format.left_indent = Cm(0.7)
                run_label = imp_p.add_run("→ Tác động: ")
                _set_font(run_label, bold=True, size=10.5, color=cfg["color"])
                run_imp = imp_p.add_run(impact)
                _set_font(run_imp, size=10.5, italic=True)

            # Nguồn & link
            src_p = doc.add_paragraph()
            src_p.paragraph_format.space_before = Pt(2)
            src_p.paragraph_format.space_after = Pt(6)
            src_p.paragraph_format.left_indent = Cm(0.7)
            run_src = src_p.add_run(f"Nguồn: {source}  |  {url}")
            _set_font(run_src, size=9, color=RGBColor(0xAA, 0xAA, 0xAA))

        doc.add_paragraph()  # khoảng cách giữa các section

    # ── Footer ────────────────────────────────────────────────────────────────
    _add_horizontal_rule(doc)
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(6)
    run_f = footer_p.add_run(
        "Bản tin tổng hợp tự động bởi AI · Chỉ mang tính tham khảo, không phải khuyến nghị đầu tư"
    )
    _set_font(run_f, size=9, color=RGBColor(0xBB, 0xBB, 0xBB), italic=True)

    # Xuất ra bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
